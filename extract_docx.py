import os
import docx
from PIL import Image
import io
from paddleocr import PaddleOCR

ocr = PaddleOCR(lang="korean")  # 최신 버전용

def extract_docx_chunks_with_images(filepath, img_dir='img_output_docx'):
    doc = docx.Document(filepath)
    os.makedirs(img_dir, exist_ok=True)
    chunks = []
    image_count = 0

    # 텍스트(문단)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append({"type": "text", "content": text})

    # 표
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append('\t'.join(row_text))
        if table_text:
            tstr = '\n'.join(table_text)
            chunks.append({"type": "table", "content": tstr})

    # 이미지 추출(+OCR)
    for i, rel in enumerate(doc.part.rels):
        rel_obj = doc.part.rels[rel]
        if "image" in rel_obj.target_ref:
            img_bytes = rel_obj.target_part.blob
            img = Image.open(io.BytesIO(img_bytes))
            img_path = os.path.join(img_dir, f"img_{i}.png")
            img.save(img_path)
            chunks.append({"type": "image", "content": img_path})

            # 최신 PaddleOCR 사용: .predict() & ['rec_texts']
            result = ocr.predict(img_path)
            ocr_lines = result[0]['rec_texts'] if result else []
            ocr_text = "\n".join(ocr_lines)
            if ocr_text:
                chunks.append({"type": "ocr", "content": ocr_text})

    return chunks

if __name__ == "__main__":
    out = extract_docx_chunks_with_images("sample.docx")
    for chunk in out:
        print(f"[{chunk['type']}]", chunk['content'][:80])
