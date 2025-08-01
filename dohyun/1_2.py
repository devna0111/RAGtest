import os
import base64
import json
from typing import List, Dict
from docx import Document
import ollama

# --------------------------
# 1. DOCX 텍스트 / 표 / 이미지 추출
# --------------------------

def extract_text_from_docx(file_path: str) -> List[Dict]:
    """DOCX에서 텍스트와 표, 이미지를 추출하는 함수"""
    doc = Document(file_path)  # DOCX 문서 열기
    chunks = []

    # 1) 텍스트 추출
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():  # 빈 문단은 제외
            chunks.append({
                "type": "text",
                "content": para.text.strip(),
                "metadata": {"position": idx+1}
            })

    # 2) 표 추출
    for t_idx, table in enumerate(doc.tables):
        data = [[cell.text for cell in row.cells] for row in table.rows]  # 표 데이터를 2차원 리스트로 변환
        chunks.append({
            "type": "table",
            "content": json.dumps(data, ensure_ascii=False),
            "metadata": {"table_index": t_idx+1}
        })

    # 3) 이미지 추출 후 qwen2.5vl로 분석
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:  # 이미지가 포함된 rel만 처리
            img_data = rel.target_part.blob  # 이미지 바이너리 데이터
            img_base64 = base64.b64encode(img_data).decode()  # Base64 인코딩
            try:
                response = ollama.chat(  # ollama로 qwen2.5vl 모델 호출
                    model="qwen2.5vl:7b",
                    messages=[{
                        "role": "user",
                        "content": "이 문서에 포함된 이미지를 설명해줘.",
                        "images": [img_base64]
                    }]
                )
                img_description = response["message"]["content"]
            except Exception as e:
                img_description = f"이미지 분석 실패: {str(e)}"

            chunks.append({
                "type": "image",
                "content": img_description,
                "metadata": {"image_index": len(chunks)+1}
            })

    return chunks


# --------------------------
# 2. 청크 분할 (텍스트 문단 길이 기준)
# --------------------------

def chunk_text(content: str, max_length=100) -> List[str]:
    """텍스트를 일정 길이로 분할"""
    words = content.split()
    return [" ".join(words[i:i+max_length]) for i in range(0, len(words), max_length)]


# --------------------------
# 3. 최종 JSON 스키마 생성
# --------------------------

def process_docx(file_path: str) -> List[Dict]:
    """DOCX 파일 전체를 처리하여 청크별 데이터 생성"""
    results = extract_text_from_docx(file_path)
    final_chunks = []

    for item in results:
        if item["type"] == "text":  # 텍스트는 분할 처리
            for idx, c in enumerate(chunk_text(item["content"], 100)):
                final_chunks.append({
                    "chunk_id": f"{item['metadata']['position']}_{idx}",
                    "type": item["type"],
                    "content": c,
                    "metadata": item["metadata"]
                })
        else:
            final_chunks.append(item)

    return final_chunks


# --------------------------
# 실행 예시
# --------------------------

if __name__ == "__main__":
    import sys
    try:
        file = "sample.docx"
        print("[INFO] DOCX 파일 처리 시작:", file)
        chunks = process_docx(file)
        print("[INFO] 추출된 청크 개수:", len(chunks))
        print("\n[INFO] 추출 결과:")
        for c in chunks:
            print(f"- TYPE: {c['type']}")
            print(f"  CONTENT: {c['content'][:200]}")  # 200자까지만 미리보기
            print("  METADATA:", c["metadata"])
            print("-" * 50)
    except Exception as e:
        print("[ERROR] 실행 중 오류 발생:", str(e))


