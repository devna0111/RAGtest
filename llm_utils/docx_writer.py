from docx import Document

def save_text_to_docx(text: str, output_path: str = "output/report.docx") -> str:
    doc = Document()
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(output_path)
    return output_path

from docx import Document
import os

def save_text_to_docx(text: str, output_path: str = "output/report.docx") -> str:
    """
    주어진 텍스트를 .docx 파일로 저장
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    sample_text = (
        "1. 개요\n"
        "이 보고서는 사내 자동화 프로젝트의 개요를 설명합니다.\n\n"
        "2. 주요 내용\n"
        "- 문서 요약 자동화\n"
        "- 발표자료 생성\n"
        "- 벡터 DB 연동\n\n"
        "3. 분석\n"
        "RAG 구조를 기반으로 다양한 문서에 대해 자동화를 시도함.\n\n"
        "4. 결론 및 제언\n"
        "문서 통합 관리 플랫폼으로의 확장을 제안함."
    )
    output_path = save_text_to_docx(sample_text, "output/sample_report.docx")
    print(f"DOCX 저장 완료: {output_path}")
